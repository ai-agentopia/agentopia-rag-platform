"""Format-aware text extraction for S3-ingested objects.

Dispatches by file extension. Designed to handle both str and bytes input
because Pathway's UDF layer delivers data as str (not bytes) when replaying
from state or when format="binary" is used with the current Pathway version.

Supported:
  .docx — python-docx paragraph extraction (bytes input only; str is warned)
  all other extensions — str returned directly, or bytes decoded as UTF-8

Note on Pathway UDF boundary behavior (0.30.0): pw.io.s3.read(format="binary")
delivers bytes to the UDF, but docx.Document() may raise BadZipFile if Pathway's
bytes representation differs from the original file bytes. The try/except below
ensures Pathway never sees an unhandled exception and silently falls back to raw
bytes — which would produce binary garbage downstream. We catch it and return "".
"""

from __future__ import annotations

import io
import logging

_log = logging.getLogger("text_extract")


def extract_text(data, path: str) -> str:
    """Return plain text for an S3 object, dispatched by file extension.

    Handles both str and bytes safely:
    - Non-DOCX str: return as-is (Pathway delivered already-decoded content)
    - Non-DOCX bytes: decode as UTF-8 with replacement chars
    - DOCX bytes: extract paragraph text via python-docx
    - DOCX str: log warning, return "" (bytes required; str recovery unproven)
    - DOCX bytes but invalid ZIP: log warning, return "" (Pathway UDF byte coercion)
    """
    ext = path.rsplit(".", 1)[-1].lower() if "." in path else ""

    if ext == "docx":
        if not isinstance(data, bytes):
            _log.warning(
                "docx extraction skipped for %s: expected bytes, got %s. "
                "Pipeline continues with empty content for this object.",
                path,
                type(data).__name__,
            )
            return ""
        try:
            import docx  # python-docx
            doc = docx.Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as exc:
            _log.warning(
                "docx extraction failed for %s: %s (%s). "
                "Pipeline continues with empty content for this object.",
                path,
                type(exc).__name__,
                exc,
            )
            return ""

    if isinstance(data, str):
        return data
    return data.decode("utf-8", errors="replace")
