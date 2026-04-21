"""Generate deterministic fixtures for each approved document type.

Each fixture contains a distinct, searchable sentinel phrase so downstream
validation can assert "this parsed chunk came from this file".
"""
import io
from pathlib import Path

import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

FIXTURE_DIR = Path(__file__).parent / "fixtures"
FIXTURE_DIR.mkdir(exist_ok=True)

# Sentinel phrases — one per type, unique
SENTINELS = {
    "docx": "DOCX sentinel alpha-architecture paragraph.",
    "pdf":  "PDF sentinel bravo-knowledge baseline.",
    "html": "HTML sentinel charlie-retrieval boundary.",
    "txt":  "TXT sentinel delta-plaintext passthrough.",
    "md":   "MD sentinel echo-markdown heading body.",
}

# --- DOCX ------------------------------------------------------------
doc = docx.Document()
doc.add_paragraph(SENTINELS["docx"])
doc.add_paragraph("Second DOCX paragraph with additional text.")
doc.save(FIXTURE_DIR / "hello.docx")

# --- PDF (text-based) -----------------------------------------------
pdf_buf = io.BytesIO()
c = canvas.Canvas(str(FIXTURE_DIR / "hello.pdf"), pagesize=letter)
c.drawString(72, 720, SENTINELS["pdf"])
c.drawString(72, 700, "Second PDF line with more content for chunk.")
c.save()

# --- HTML -----------------------------------------------------------
(FIXTURE_DIR / "hello.html").write_text(
    f"<html><head><title>t</title></head>"
    f"<body><script>evil()</script><p>{SENTINELS['html']}</p>"
    f"<p>Second HTML paragraph.</p></body></html>",
    encoding="utf-8",
)

# --- TXT ------------------------------------------------------------
(FIXTURE_DIR / "hello.txt").write_text(
    f"{SENTINELS['txt']}\nSecond line of plain text.", encoding="utf-8"
)

# --- MD -------------------------------------------------------------
(FIXTURE_DIR / "hello.md").write_text(
    f"# Heading\n\n{SENTINELS['md']}\n\n## Sub\n\nSecond MD paragraph.",
    encoding="utf-8",
)

print("Generated fixtures:")
for f in sorted(FIXTURE_DIR.iterdir()):
    if f.is_file():
        print(f"  {f.name}  ({f.stat().st_size} bytes)")
